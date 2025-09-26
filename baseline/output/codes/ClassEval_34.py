from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        doc = Document(self.file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        doc = Document()
        p = doc.add_paragraph()
        p.alignment = alignment
        p.style.font.size = Pt(font_size)
        p.add_run(content)
        doc.save(self.file_path)
        return True

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        doc = Document()
        p = doc.add_heading(heading, level)
        doc.save(self.file_path)
        return True

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        doc = Document()
        table = doc.add_table(rows=1, cols=len(data[0]))
        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
        doc.save(self.file_path)
        return True

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT

import unittest
import os


class DocFileHandlerTestReadText(unittest.TestCase):
    def test_read_text_1(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "Initial content"
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_read_text_2(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("111")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "111"
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_read_text_3(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("aaa")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "aaa"
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_read_text_4(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("aaa\nbbb")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "aaa\nbbb"
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_read_text_5(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = ""
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)


class DocFileHandlerTestWriteText(unittest.TestCase):
    def setUp(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

    def tearDown(self):
        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_write_text_1(self):
        new_content = "New content 1"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

    def test_write_text_2(self):
        new_content = "New content 2"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

    def test_write_text_3(self):
        new_content = "New content 3"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

    def test_write_text_4(self):
        new_content = "New content 4"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

    def test_write_text_5(self):
        new_content = ""
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)


class DocFileHandlerTestAddHeading(unittest.TestCase):
    def setUp(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

    def tearDown(self):
        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_add_heading_1(self):
        heading = "Test Heading 1"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

    def test_add_heading_2(self):
        heading = "Test Heading 2"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

    def test_add_heading_3(self):
        heading = "Test Heading 3"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

    def test_add_heading_4(self):
        heading = "Test Heading 4"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

    def test_add_heading_5(self):
        heading = "Test Heading 5"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)


class DocFileHandlerTestAddTable(unittest.TestCase):
    def setUp(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

    def tearDown(self):
        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_add_table_1(self):
        data = [['Name', 'Age']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)
        self.assertEqual(len(table.columns), 2)

    def test_add_table_2(self):
        data = [['Name', 'Age'], ['John', '25']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, 'John')

    def test_add_table_3(self):
        data = [['Name', 'Age'], ['John', '25'], ['Emma', '30']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, 'John')
        self.assertEqual(table.cell(2, 1).text, '30')

    def test_add_table_4(self):
        data = [['Name', 'Age'], ['aaa', '25'], ['Emma', '30']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, 'aaa')
        self.assertEqual(table.cell(2, 1).text, '30')

    def test_add_table_5(self):
        data = [['Name', 'Age'], ['John', '25'], ['Emma', '90']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, 'John')
        self.assertEqual(table.cell(2, 1).text, '90')


class DocFileHandlerTest(unittest.TestCase):
    def test_DocFileHandler(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "Initial content"
        self.assertEqual(text_content, expected_content)

        new_content = "New content 1"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

        heading = "Test Heading 1"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

        data = [['Name', 'Age']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)
        self.assertEqual(len(table.columns), 2)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)